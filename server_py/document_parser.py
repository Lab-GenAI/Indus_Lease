import os
import subprocess
import tempfile
import uuid
import shutil

SUPPORTED_EXTENSIONS = ["pdf", "docx", "eml", "msg", "txt"]
ATTACHMENT_EXTENSIONS = {"pdf", "docx", "txt", "eml", "msg","xlsx","xls"}


def get_file_extension(file_name: str) -> str:
    _, ext = os.path.splitext(file_name)
    return ext.lower().lstrip(".")


def is_supported_file(file_name: str) -> bool:
    return get_file_extension(file_name) in SUPPORTED_EXTENSIONS


def parse_document(file_path: str, file_type: str) -> str:
    try:
        ext = file_type.lower()
        if ext == "pdf":
            return _parse_pdf(file_path)
        elif ext == "docx":
            return _parse_docx(file_path)
        elif ext == "eml":
            return _parse_eml(file_path)
        elif ext == "msg":
            return _parse_msg(file_path)
        elif ext == "txt":
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        elif ext in {"xlsx","xls"}:
            return _parse_excel(file_path)
        else:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
    except Exception as e:
        print(f"Error parsing {file_path}: {e}")
        return f"[Error parsing file: {e}]"


SCANNER_WATERMARKS = [
    "scanned by camscanner",
    "scanned by tapscanner",
    "scanned by adobe scan",
    "scanned by microsoft lens",
    "scanned by genius scan",
    "created by camscanner",
    "scanned document",
]


def _is_meaningful_text(text: str) -> bool:
    import re
    stripped = text.replace(" ", "").replace("\n", "")
    if len(stripped) <= 50:
        return False
    lower_text = text.lower().strip()
    cleaned = re.sub(r'\s+', ' ', lower_text).strip()
    for wm in SCANNER_WATERMARKS:
        without_wm = cleaned.replace(wm, "").strip()
        without_wm = re.sub(r'\s+', ' ', without_wm).strip()
        if len(without_wm) < 20:
            print(f"[PDF] Detected scanner watermark only: '{wm}', triggering OCR")
            return False
    unique_words = set(re.findall(r'[a-z]{3,}', lower_text))
    if len(unique_words) < 5 and len(stripped) < 200:
        print(f"[PDF] Text too repetitive ({len(unique_words)} unique words), triggering OCR")
        return False
    return True


def _parse_pdf(file_path: str) -> str:
    try:
        import pdfplumber
        text_parts = []
        table_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                try:
                    tables = page.extract_tables()
                    for table in tables:
                        if not table:
                            continue
                        table_rows = []
                        for row in table:
                            if row:
                                cells = [str(cell).strip() if cell else "" for cell in row]
                                if any(cells):
                                    table_rows.append(" | ".join(cells))
                        if table_rows:
                            table_parts.append(f"[Table on page {page_num + 1}]\n" + "\n".join(table_rows))
                except Exception:
                    pass
        text = "\n".join(text_parts).strip()
        if table_parts:
            text += "\n\n--- TABLES ---\n\n" + "\n\n".join(table_parts)
        if _is_meaningful_text(text):
            return text
        print(f"PDF appears to be scanned, attempting OCR: {file_path}")
        return _ocr_pdf(file_path)
    except Exception as e:
        print(f"PDF parse error, attempting OCR: {e}")
        try:
            return _ocr_pdf(file_path)
        except Exception as ocr_e:
            return f"[PDF parsing and OCR failed: {e}]"



def _clean_ocr_text(text: str) -> str:
    import re
    text = re.sub(r'[^\S\n]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if len(line) <= 1 and not line.isalnum():
            continue
        if len(line) > 2 and sum(1 for c in line if not c.isalnum() and c not in ' .,;:!?-/()@#$%&*"\'+') / len(line) > 0.5:
            continue
        cleaned_lines.append(line)
    return '\n'.join(cleaned_lines).strip()


def _ocr_pdf(pdf_path: str) -> str:
    temp_dir = os.path.join(tempfile.gettempdir(), f"ocr-{uuid.uuid4().hex}")
    os.makedirs(temp_dir, exist_ok=True)

    try:
        poppler_path = os.environ.get("POPPLER_PATH")
        pdftoppm_bin = "pdftoppm"
        if poppler_path:
            candidate = os.path.join(poppler_path, "pdftoppm")
            if os.path.exists(candidate):
                pdftoppm_bin = candidate
            candidate_exe = os.path.join(poppler_path, "pdftoppm.exe")
            if os.path.exists(candidate_exe):
                pdftoppm_bin = candidate_exe

        page_prefix = os.path.join(temp_dir, "page")
        result = subprocess.run(
            [pdftoppm_bin, "-png", "-r", "300", pdf_path, page_prefix],
            capture_output=True, text=True, timeout=180
        )
        if result.returncode != 0:
            print(f"[OCR] pdftoppm failed (code {result.returncode}): {result.stderr}")
            return f"[OCR failed: pdftoppm exited with code {result.returncode}]"

        image_files = sorted([
            os.path.join(temp_dir, f) for f in os.listdir(temp_dir) if f.endswith(".png")
        ])
        if not image_files:
            return "[Scanned PDF: No pages could be converted for OCR]"

        print(f"[OCR] Converted {len(image_files)} pages at 300 DPI, sending to Vision API...")

        page_texts = []
        for i, img_file in enumerate(image_files):
            page_text = _ocr_page_with_vision(img_file, i)
            if page_text:
                page_texts.append(f"--- Page {i + 1} ---\n{page_text}")
                print(f"[OCR] Page {i + 1}: extracted {len(page_text)} chars via Vision")
            else:
                print(f"[OCR] Page {i + 1}: no text extracted via Vision")

        if not page_texts:
            return "[Scanned PDF: Vision OCR produced no text from any page]"

        full_text = "\n\n".join(page_texts)
        result_text = _clean_ocr_text(full_text)
        print(f"[OCR] Total: {len(result_text)} chars from {len(page_texts)}/{len(image_files)} pages")
        return result_text if len(result_text) > 10 else "[Scanned PDF: OCR produced minimal text]"

    except Exception as e:
        print(f"[OCR] Error: {e}")
        return f"[OCR failed: {e}]"
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def _get_vision_client_and_model():
    from server_py.config import get_config
    from openai import OpenAI
    config = get_config()
    model = config.get("extraction_model", "gpt-4.1")
    if model.startswith("claude-"):
        model = "gpt-4.1"
    api_key = config.get("openai_api_key", "")
    base_url = config.get("openai_base_url", "https://api.openai.com/v1")
    client = OpenAI(api_key=api_key, base_url=base_url)
    return client, model


def _ocr_page_with_vision(image_path: str, page_index: int) -> str:
    import base64
    import time
    from server_py.cost_tracker import log_cost

    client, ocr_model = _get_vision_client_and_model()
    MAX_RETRIES = 3

    try:
        with open(image_path, "rb") as f:
            image_data = base64.b64encode(f.read()).decode("utf-8")
    except Exception as e:
        print(f"[OCR] Failed to read image file for page {page_index + 1}: {e}")
        return ""

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            response = client.chat.completions.create(
                model=ocr_model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Extract ALL text from this scanned document page. Preserve the original layout, formatting, paragraph structure, and table structure as much as possible. Include all headers, footers, page numbers, stamps, and any text visible on the page. Output ONLY the extracted text, nothing else.",
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{image_data}",
                                    "detail": "high",
                                },
                            },
                        ],
                    }
                ],
                max_completion_tokens=4096,
                timeout=120,
            )

            usage = response.usage
            if usage:
                log_cost(
                    type_="vision_ocr",
                    model=ocr_model,
                    input_tokens=usage.prompt_tokens,
                    output_tokens=usage.completion_tokens,
                )

            text = (response.choices[0].message.content or "").strip()
            return text
        except Exception as e:
            if attempt < MAX_RETRIES:
                wait_time = 2 ** attempt
                print(f"[OCR] Vision API error on page {page_index + 1} (attempt {attempt}/{MAX_RETRIES}): {e} — retrying in {wait_time}s...")
                time.sleep(wait_time)
            else:
                print(f"[OCR] Vision API error on page {page_index + 1} (attempt {attempt}/{MAX_RETRIES}): {e} — giving up")
                return ""


def _parse_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        table_parts = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                table_parts.append(f"[Table {i + 1}]\n" + "\n".join(rows))
        text = "\n".join(parts)
        if table_parts:
            text += "\n\n--- TABLES ---\n\n" + "\n\n".join(table_parts)
        return text if text.strip() else "[Empty DOCX document]"
    except Exception as e:
        return f"[DOCX parsing failed: {e}]"


def _parse_eml(file_path: str) -> str:
    try:
        import email
        from email import policy
        with open(file_path, "rb") as f:
            msg = email.message_from_binary_file(f, policy=policy.default)
        parts = []
        if msg["subject"]:
            parts.append(f"Subject: {msg['subject']}")
        if msg["from"]:
            parts.append(f"From: {msg['from']}")
        if msg["to"]:
            parts.append(f"To: {msg['to']}")
        if msg.get("cc"):
            parts.append(f"Cc: {msg['cc']}")
        if msg["date"]:
            parts.append(f"Date: {msg['date']}")

        body_text = ""
        plain_body = msg.get_body(preferencelist=("plain",))
        if plain_body:
            content = plain_body.get_content()
            if content and content.strip():
                body_text = content.strip()

        if not body_text or len(body_text) < 20:
            html_body = msg.get_body(preferencelist=("html",))
            if html_body:
                html_content = html_body.get_content()
                if html_content:
                    converted = _strip_html_to_text(html_content)
                    if len(converted) > len(body_text):
                        body_text = converted

        if body_text:
            parts.append(f"\nBody:\n{body_text}")

        return "\n".join(parts) if parts else "[Empty email]"
    except Exception as e:
        return f"[EML parsing failed: {e}]"


def extract_email_attachments(file_path: str, file_type: str) -> list:
    """Extract attachments from EML/MSG files, save to temp dir, return list of dicts with file_name, file_path, file_type."""
    attachments = []
    temp_dir = os.path.join(tempfile.gettempdir(), f"email-att-{uuid.uuid4().hex}")

    try:
        if file_type == "eml":
            attachments = _extract_eml_attachments(file_path, temp_dir)
        elif file_type == "msg":
            attachments = _extract_msg_attachments(file_path, temp_dir)
    except Exception as e:
        print(f"[ATTACHMENT] Error extracting attachments from {file_path}: {e}")

    if not attachments and os.path.exists(temp_dir):
        shutil.rmtree(temp_dir, ignore_errors=True)

    return attachments


def _extract_eml_attachments(file_path: str, temp_dir: str) -> list:
    import email
    from email import policy

    with open(file_path, "rb") as f:
        msg = email.message_from_binary_file(f, policy=policy.default)

    attachments = []
    for part in msg.walk():
        content_disposition = str(part.get("Content-Disposition", ""))
        if "attachment" not in content_disposition and "inline" not in content_disposition:
            continue

        filename = part.get_filename()
        if not filename:
            continue

        ext = get_file_extension(filename)
        if ext not in ATTACHMENT_EXTENSIONS:
            print(f"[ATTACHMENT] Skipping unsupported attachment: {filename} (.{ext})")
            continue

        os.makedirs(temp_dir, exist_ok=True)
        safe_name = f"{uuid.uuid4().hex}_{filename}"
        att_path = os.path.join(temp_dir, safe_name)

        payload = part.get_payload(decode=True)
        if payload:
            with open(att_path, "wb") as af:
                af.write(payload)
            attachments.append({
                "file_name": filename,
                "file_path": att_path,
                "file_type": ext,
                "source": os.path.basename(file_path),
            })
            print(f"[ATTACHMENT] Extracted from EML: {filename} ({len(payload):,} bytes)")

    return attachments


def _extract_msg_attachments(file_path: str, temp_dir: str) -> list:
    try:
        import extract_msg
    except ImportError:
        print("[ATTACHMENT] extract_msg not installed, cannot extract MSG attachments")
        return []

    msg = extract_msg.Message(file_path)
    attachments = []

    try:
        if not msg.attachments:
            return []

        for att in msg.attachments:
            filename = att.longFilename or att.shortFilename or "unnamed"
            ext = get_file_extension(filename)
            if ext not in ATTACHMENT_EXTENSIONS:
                print(f"[ATTACHMENT] Skipping unsupported attachment: {filename} (.{ext})")
                continue

            os.makedirs(temp_dir, exist_ok=True)
            safe_name = f"{uuid.uuid4().hex}_{filename}"
            att_path = os.path.join(temp_dir, safe_name)

            try:
                data = att.data
                if data:
                    with open(att_path, "wb") as af:
                        af.write(data)
                    attachments.append({
                        "file_name": filename,
                        "file_path": att_path,
                        "file_type": ext,
                        "source": os.path.basename(file_path),
                    })
                    print(f"[ATTACHMENT] Extracted from MSG: {filename} ({len(data):,} bytes)")
            except Exception as e:
                print(f"[ATTACHMENT] Failed to extract {filename}: {e}")
    finally:
        msg.close()

    return attachments


def _strip_html_to_text(html_content: str) -> str:
    import re
    text = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</(p|div|tr|li|h[1-6])>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<td[^>]*>', ' | ', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'&nbsp;', ' ', text)
    text = re.sub(r'&amp;', '&', text)
    text = re.sub(r'&lt;', '<', text)
    text = re.sub(r'&gt;', '>', text)
    text = re.sub(r'&quot;', '"', text)
    text = re.sub(r'&#\d+;', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _parse_msg(file_path: str) -> str:
    try:
        import extract_msg
        msg = extract_msg.Message(file_path)
        parts = []
        if msg.sender:
            parts.append(f"From: {msg.sender}")
        if msg.to:
            parts.append(f"To: {msg.to}")
        try:
            if msg.cc:
                parts.append(f"Cc: {msg.cc}")
        except Exception:
            pass
        if msg.subject:
            parts.append(f"Subject: {msg.subject}")
        if msg.date:
            parts.append(f"Date: {msg.date}")
        if parts:
            parts.append("---")

        body_text = ""
        if msg.body:
            body_text = msg.body
        if not body_text or len(body_text.strip()) < 20:
            try:
                html_body = msg.htmlBody
                if html_body:
                    if isinstance(html_body, bytes):
                        html_body = html_body.decode("utf-8", errors="replace")
                    converted = _strip_html_to_text(html_body)
                    if len(converted) > len(body_text or ""):
                        body_text = converted
                        print(f"[MSG] Used HTML body (converted {len(converted)} chars)")
            except Exception as e:
                print(f"[MSG] HTML body extraction failed: {e}")

        if body_text:
            parts.append(body_text)
        else:
            parts.append("[No text body found in MSG file]")

        msg.close()
        return "\n".join(parts)
    except ImportError:
        try:
            with open(file_path, "rb") as f:
                data = f.read()
            text = data.decode("utf-8", errors="replace")
            printable = "".join(c if c.isprintable() or c in "\n\r\t" else " " for c in text)
            cleaned = "\n".join(line.strip() for line in printable.split("\n") if line.strip())
            return cleaned if len(cleaned) > 20 else "[MSG file: minimal text extracted]"
        except Exception as e:
            return f"[MSG file parsing failed: {e}]"
    except Exception as e:
        return f"[MSG file parsing failed: {e}]"

def _parse_excel(file_path: str) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append(" | ".join(cells))
            if rows:
                parts.append(f"=== Sheet: {sheet_name} ===")
                parts.extend(rows)
        wb.close()
        return "\n".join(parts) if parts else "[Empty Excel file]"
    except Exception as e:
        return f"[Excel parsing failed: {e}]"

def render_email_as_html(file_path: str, ext: str) -> str:
    if ext == "eml":
        try:
            import email
            from email import policy
            with open(file_path, "rb") as f:
                msg = email.message_from_binary_file(f, policy=policy.default)
            subject = msg.get("subject", "")
            from_addr = msg.get("from", "")
            to_addr = msg.get("to", "")
            date = msg.get("date", "")
            body_part = msg.get_body(preferencelist=("html", "plain"))
            body_html = ""
            body_text = ""
            if body_part:
                ct = body_part.get_content_type()
                content = body_part.get_content()
                if "html" in ct:
                    body_html = content
                else:
                    body_text = content
            return _build_email_html(from_addr, to_addr, "", subject, date, body_text, body_html, [])
        except Exception as e:
            return _build_email_html("", "", "", "Error", "", f"Failed to parse EML: {e}", "", [])

    if ext == "msg":
        try:
            import extract_msg
            msg = extract_msg.Message(file_path)
            attachment_names = []
            if msg.attachments:
                attachment_names = [att.longFilename or att.shortFilename or "unnamed" for att in msg.attachments]
            html = _build_email_html(
                str(msg.sender or ""), str(msg.to or ""), str(msg.cc or ""),
                str(msg.subject or ""), str(msg.date or ""),
                str(msg.body or ""), "", attachment_names
            )
            msg.close()
            return html
        except Exception as e:
            return _build_email_html("", "", "", "Error", "", f"Failed to parse MSG: {e}", "", [])

    return _build_email_html("", "", "", "Unsupported", "", "This file type cannot be previewed as email.", "", [])


def _escape_html(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _build_email_html(from_addr, to_addr, cc, subject, date, body_text, body_html, attachments):
    header_rows = []
    if from_addr:
        header_rows.append(f'<tr><td class="label">From</td><td>{_escape_html(str(from_addr))}</td></tr>')
    if to_addr:
        header_rows.append(f'<tr><td class="label">To</td><td>{_escape_html(str(to_addr))}</td></tr>')
    if cc:
        header_rows.append(f'<tr><td class="label">Cc</td><td>{_escape_html(str(cc))}</td></tr>')
    if subject:
        header_rows.append(f'<tr><td class="label">Subject</td><td><strong>{_escape_html(str(subject))}</strong></td></tr>')
    if date:
        header_rows.append(f'<tr><td class="label">Date</td><td>{_escape_html(str(date))}</td></tr>')

    if body_html:
        body_content = body_html
    elif body_text:
        body_content = f'<pre style="white-space:pre-wrap;word-wrap:break-word;font-family:inherit;margin:0;">{_escape_html(str(body_text))}</pre>'
    else:
        body_content = '<p style="color:#888;">No message body.</p>'

    attachment_section = ""
    if attachments:
        att_list = ", ".join(_escape_html(a) for a in attachments)
        attachment_section = f'<div class="attachments"><span class="label">Attachments:</span> {att_list}</div>'

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  * {{ box-sizing: border-box; }}
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; margin: 0; padding: 24px; background: #1a1a2e; color: #e0e0e0; line-height: 1.6; }}
  .header {{ background: #16213e; border: 1px solid #2a2a4a; border-radius: 8px; padding: 16px; margin-bottom: 20px; }}
  .header table {{ width: 100%; border-collapse: collapse; }}
  .header td {{ padding: 6px 12px; vertical-align: top; font-size: 14px; }}
  .header td.label {{ color: #888; font-weight: 600; width: 80px; white-space: nowrap; }}
  .body-content {{ background: #16213e; border: 1px solid #2a2a4a; border-radius: 8px; padding: 20px; font-size: 14px; }}
  .attachments {{ margin-top: 12px; padding: 10px 12px; background: #1a1a3e; border-radius: 6px; font-size: 13px; color: #aaa; }}
  .attachments .label {{ font-weight: 600; }}
  a {{ color: #6ea8fe; }}
  pre {{ font-size: 14px; }}
</style>
</head>
<body>
  <div class="header">
    <table>{"".join(header_rows)}</table>
    {attachment_section}
  </div>
  <div class="body-content">{body_content}</div>
</body>
</html>"""
